"""
proposal_builder.py

Fully Populated & Structured Proposal Builder.
Integrates Phase 6 (AI Generated Answers) directly into Phase 7 (Proposal Outline)
and exports an enterprise-ready Microsoft Word (.docx) document without empty placeholders.
"""

import re
from typing import List, Dict, Any, Optional
from types import SimpleNamespace
import io
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from schemas import ProposalOutline, GeneratedResponse, Deliverable, KeyDatesBudget


# --- Helper: Set Table Cell Background Color ---
def set_cell_background(cell, fill_hex: str):
    """Sets background color of a table cell in docx."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


# ---------------------------------------------------------------------------
# Question-to-section matching
#
# PREVIOUS APPROACH (removed): a hand-maintained keyword dict where a
# category (e.g. "pricing" -> ["price","cost","financial",...]) only even
# got CHECKED if the category name literally appeared as a substring in the
# section's own title (`if rule_key in child_title_lower`). Real outline
# titles are AI-generated and specific ("Unified Search Experience",
# "AI-Enhanced Search Capabilities") — they essentially never contain a
# generic category word like "approach" or "scope" verbatim, so most
# sections never even reached the keyword check and fell through to the
# generic boilerplate fallback. Worse, a single overly-generic keyword
# (e.g. "schedule" under "pricing") could false-positive match a completely
# unrelated question ("What training and support will you provide
# post-implementation?" landed in "Pricing Structure" this way).
#
# NEW APPROACH: score every (question, section) pair by how many
# significant words they actually share, and assign each question to
# whichever section scores highest — no reliance on a category name
# appearing literally anywhere. This is directly tied to the real words in
# the real title/question rather than a keyword list that can never cover
# every phrasing. A question only ever lands in ONE section (its best
# match), and anything that scores 0 everywhere (no shared words at all)
# is NOT silently dropped — it's collected into a final "Additional
# Requirements & Responses" appendix, so a low-confidence or unmatchable
# question still makes it into the document instead of vanishing.
# ---------------------------------------------------------------------------
_STOPWORDS = {
    "the", "a", "an", "and", "or", "of", "to", "in", "on", "for", "with", "will",
    "you", "your", "what", "how", "why", "when", "where", "describe", "provide",
    "is", "are", "be", "this", "that", "it", "as", "by", "we", "our", "their",
    "they", "proposal", "response", "please", "explain", "which", "do", "does",
    "should", "would", "any", "all", "each", "if", "not", "have", "has", "can",
}


def _significant_words(text: str) -> set:
    words = re.findall(r"[a-z]+", (text or "").lower())
    # >= 2 chars (not > 2) so short domain terms like "AI", "ML", "IT" survive —
    # every short word that would otherwise be noise (an, by, to, in, on, is,
    # be, it, as, we, ...) is already covered by _STOPWORDS above.
    return {w for w in words if len(w) >= 2 and w not in _STOPWORDS}


def _match_score(question_words: set, title_words: set) -> int:
    return len(question_words & title_words)


# --- 1. Assembly Engine: Direct Mapping & Population ---
def assemble_proposal_content(
    outline: ProposalOutline,
    generated_responses: List[GeneratedResponse],
    rfp_identifier: Optional[str] = None,
    company_name: str = "SPS"
) -> List[Dict[str, Any]]:
    """
    Merges AI responses into matching outline sections based on shared
    significant words between each question and each section title.
    Eliminates empty placeholder tags by generating context-aware fallback text.
    Anything that can't be confidently matched anywhere lands in a final
    "Additional Requirements & Responses" appendix rather than being lost.
    """
    # --- Pass 1: score every question against every (section, child), and
    # assign each question to its single best-scoring match (score must be
    # > 0 — no shared words means no assignment, not a forced guess). ---
    flat_children = []  # [(sec_idx, subsec_idx, section_title, child), ...] in outline order
    sec_idx = 1
    for sec in outline.sections:
        subsec_idx = 1
        for child in sec.children:
            flat_children.append((sec_idx, subsec_idx, sec.title, child))
            subsec_idx += 1
        sec_idx += 1

    title_word_cache = {
        id(child): _significant_words(child.title) for *_ , child in flat_children
    }

    assignments = {id(child): [] for *_, child in flat_children}  # id(child) -> [resp, ...]
    unmatched_responses = []

    for resp in generated_responses:
        # Scored on the question's own text ONLY — NOT resp.basedOn. basedOn
        # is a knowledge-base *category* label (e.g. "security_compliance"),
        # not real RFP/outline language, and including it caused its own
        # false positive during testing: a training/support question tagged
        # basedOn="security_compliance" matched into an "E-Verify Compliance"
        # section purely because both happen to contain the word
        # "compliance" — completely unrelated to what the question actually
        # asks. The question text alone is a far more reliable signal.
        q_words = _significant_words(resp.question)
        best_child = None
        best_score = 0
        for *_, child in flat_children:
            score = _match_score(q_words, title_word_cache[id(child)])
            if score > best_score:
                best_score = score
                best_child = child
        if best_child is not None:
            assignments[id(best_child)].append(resp)
        else:
            unmatched_responses.append(resp)

    # --- Pass 2: build assembled_sections using those assignments ---
    assembled_sections = []
    sec_idx = 1
    for sec in outline.sections:
        section_number = f"{sec_idx}.0"
        sub_sections = []

        subsec_idx = 1
        for child in sec.children:
            child_number = f"{sec_idx}.{subsec_idx}"
            child_title_lower = child.title.lower()
            matched_answers = assignments[id(child)]

            # Build narrative text for this subsection
            if matched_answers:
                content_blocks = []
                for ans in matched_answers:
                    content_blocks.append(
                        f"**Requirement / Question:** {ans.question}\n\n"
                        f"**{company_name} Proposed Solution:**\n{ans.response}"
                    )
                content_text = "\n\n".join(content_blocks)
            else:
                # Dynamic Smart Fallbacks for Standard Admin/Structural Sections
                if "cover" in child_title_lower:
                    content_text = f"Official Proposal Document submitted by {company_name} in response to RFP {rfp_identifier or 'Requirement'}."
                elif "transmittal" in child_title_lower:
                    content_text = (
                        f"This proposal constitutes a formal and binding offer by {company_name} to fulfill all scope of work, "
                        f"technical specifications, and operational SLAs outlined in RFP {rfp_identifier or ''}. "
                        f"{company_name} confirms full compliance with all administrative and legal terms."
                    )
                elif any(k in child_title_lower for k in ["affidavit", "attachment", "form"]):
                    content_text = (
                        f"The completed, signed, and notarized {child.title} has been executed by an authorized representative "
                        f"of {company_name} and is included in the final submission package index."
                    )
                elif any(k in child_title_lower for k in ["price", "cost", "pricing"]):
                    content_text = (
                        f"The financial proposal and total cost matrix for {company_name} are prepared in strict accordance "
                        f"with the client pricing guidelines. Detailed cost breakdowns and payment schedules are attached in Section 2."
                    )
                elif "reference" in child_title_lower:
                    content_text = (
                        f"{company_name} maintains a proven track record of delivering enterprise IAM and cloud infrastructure projects. "
                        f"Detailed client contact references and verified project outcomes are available upon request."
                    )
                else:
                    content_text = (
                        f"{company_name} fully acknowledges and accepts the requirements detailed under '{child.title}'. "
                        f"Our operational standard operating procedures (SOPs) ensure high availability, security, and full execution quality."
                    )

            sub_sections.append({
                "number": child_number,
                "title": child.title,
                "content": content_text,
                "qa_pairs": matched_answers
            })
            subsec_idx += 1

        assembled_sections.append({
            "number": section_number,
            "title": sec.title,
            "sub_sections": sub_sections
        })
        sec_idx += 1

    # --- Appendix: anything that shared zero significant words with EVERY
    # outline section title. Rather than silently dropping these (or, worse,
    # force-matching them somewhere wrong the way the old keyword-substring
    # approach did), they get one final catch-all section so every drafted
    # answer always makes it into the document, even if a human still needs
    # to manually move it to the right spot. ---
    if unmatched_responses:
        content_blocks = [
            f"**Requirement / Question:** {ans.question}\n\n"
            f"**{company_name} Proposed Solution:**\n{ans.response}"
            for ans in unmatched_responses
        ]
        assembled_sections.append({
            "number": f"{sec_idx}.0",
            "title": "Additional Requirements & Responses",
            "sub_sections": [{
                "number": f"{sec_idx}.1",
                "title": "Drafted Answers Not Matched to a Specific Outline Section",
                "content": "\n\n".join(content_blocks),
                "qa_pairs": unmatched_responses,
            }],
        })

    return assembled_sections


# --- 2. Professional Word Exporter Engine ---
def export_proposal_to_docx(
    rfp_identifier: str,
    assembled_sections: List[Dict[str, Any]],
    deliverables: List[Deliverable],
    dates_budget: KeyDatesBudget,
    company_name: str = "SPS"
) -> bytes:
    """
    Generates a fully populated, styled Microsoft Word (.docx) proposal file.
    """
    doc = Document()

    # --- Page Setup & Margins ---
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # --- Typography & Colors ---
    PRIMARY_COLOR = RGBColor(15, 23, 42)     # Deep Slate
    SECONDARY_COLOR = RGBColor(14, 116, 144) # Teal Accent
    TEXT_COLOR = RGBColor(51, 65, 85)        # Slate Text Body

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = TEXT_COLOR
    normal_style.paragraph_format.line_spacing = 1.2
    normal_style.paragraph_format.space_after = Pt(6)

    # --- 1. COVER PAGE ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(100)
    r_title = p_title.add_run("TECHNICAL & COMMERCIAL PROPOSAL")
    r_title.font.size = Pt(24)
    r_title.font.bold = True
    r_title.font.color.rgb = PRIMARY_COLOR

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(150)
    r_sub = p_sub.add_run(f"In Response to RFP: {rfp_identifier or 'BPM057272'}")
    r_sub.font.size = Pt(14)
    r_sub.font.color.rgb = SECONDARY_COLOR

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = p_meta.add_run(
        f"Submitted By: {company_name}\n"
        f"Submission Deadline: {dates_budget.submissionDeadline or 'As Specified in RFP'}"
    )
    r_meta.font.size = Pt(10.5)
    r_meta.font.italic = True

    doc.add_page_break()

    # --- 2. DELIVERABLES SUMMARY TABLE ---
    if deliverables:
        h_deliv = doc.add_heading("Project Key Deliverables", level=1)
        h_deliv.runs[0].font.color.rgb = PRIMARY_COLOR
        h_deliv.runs[0].font.size = Pt(16)
        
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        hdr_cells = table.rows[0].cells
        headers = ["Deliverable Description", "Mandatory", "Priority", "Est. Effort (Days)"]
        widths = [Inches(3.3), Inches(0.9), Inches(0.9), Inches(1.4)]

        for i, header_text in enumerate(headers):
            hdr_cells[i].text = header_text
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            set_cell_background(hdr_cells[i], "0F172A")
            hdr_cells[i].width = widths[i]

        for d in deliverables:
            row_cells = table.add_row().cells
            row_cells[0].text = d.description
            row_cells[1].text = "Yes" if d.mandatory else "Optional"
            row_cells[2].text = str(d.priority).upper()
            row_cells[3].text = str(d.estimatedDays or "TBD")

            for i, w in enumerate(widths):
                row_cells[i].width = w

        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- 3. FULL POPULATED PROPOSAL SECTIONS ---
    for sec in assembled_sections:
        h_sec = doc.add_heading(f"{sec['number']} {sec['title']}", level=1)
        h_sec.runs[0].font.color.rgb = PRIMARY_COLOR
        h_sec.runs[0].font.size = Pt(15)
        h_sec.paragraph_format.space_before = Pt(16)

        for subsec in sec["sub_sections"]:
            h_sub = doc.add_heading(f"{subsec['number']} {subsec['title']}", level=2)
            h_sub.runs[0].font.color.rgb = SECONDARY_COLOR
            h_sub.runs[0].font.size = Pt(12.5)
            h_sub.paragraph_format.space_before = Pt(10)

            # Insert Narrative Paragraphs
            content = subsec["content"]
            lines = content.split("\n\n")
            for line in lines:
                p = doc.add_paragraph()
                if line.startswith("**") and "**" in line[2:]:
                    # Handle Bold Headers within Text
                    parts = line.split("**")
                    for idx, part in enumerate(parts):
                        run = p.add_run(part)
                        if idx % 2 == 1:
                            run.bold = True
                            run.font.color.rgb = PRIMARY_COLOR
                else:
                    p.add_run(line)

    # --- Save to Bytes ---
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# --- 3. Compatibility bridge: build_proposal_docx() -------------------------
# ---------------------------------------------------------------------------
# app.py (and history_store.py's saved records) work with the analysis
# result as PLAIN dicts/lists throughout — analysis["proposalOutline"] is a
# dict, st.session_state.generated_responses is a list of
# {"question","response","basedOn"} dicts, etc. There's never a point in the
# app where those get turned into actual schemas.py Pydantic instances.
#
# assemble_proposal_content()/export_proposal_to_docx() above expect dotted
# attribute access (outline.sections, sec.children, child.title,
# resp.question, d.mandatory, dates_budget.submissionDeadline) — which plain
# dicts don't support (dict.sections raises AttributeError, not a lookup).
#
# Rather than parsing through the strict schemas.py Pydantic models here
# (which would crash proposal export on perfectly normal but slightly sparse
# AI output — e.g. Deliverable requires at least one `points` entry), this
# wraps each dict in a plain SimpleNamespace: enough to satisfy attribute
# access, with sane defaults filled in for anything missing, so export
# degrades gracefully instead of hard-failing.
#
# build_proposal_docx() is the function app.py actually imports and calls —
# kept under this name/signature so app.py doesn't need to change how it
# calls into this module.

def _ns(d: dict, defaults: dict) -> SimpleNamespace:
    merged = dict(defaults)
    merged.update(d or {})
    return SimpleNamespace(**merged)


def _outline_from_dict(outline: Any) -> SimpleNamespace:
    outline_dict = outline if isinstance(outline, dict) else {}
    sections = []
    for sec in outline_dict.get("sections", []) or []:
        children = [_ns(c, {"title": ""}) for c in (sec.get("children") or [])]
        sections.append(_ns({"title": sec.get("title", ""), "children": children}, {}))
    return SimpleNamespace(sections=sections)


def _responses_from_list(responses: Optional[list]) -> List[SimpleNamespace]:
    return [_ns(r, {"question": "", "response": "", "basedOn": None}) for r in (responses or [])]


def _deliverables_from_list(deliverables: Optional[list]) -> List[SimpleNamespace]:
    return [
        _ns(d, {"description": "", "mandatory": False, "priority": "Medium", "estimatedDays": None})
        for d in (deliverables or [])
    ]


def _dates_budget_from_dict(dates_budget: Optional[dict]) -> SimpleNamespace:
    return _ns(dates_budget or {}, {"submissionDeadline": None})


def build_proposal_docx(
    rfp_identifier: str,
    company_name: str,
    outline: dict,
    generated_responses: list,
    source_label: str = "",
    deliverables: Optional[list] = None,
    dates_budget: Optional[dict] = None,
) -> bytes:
    """
    The entry point app.py calls. Accepts everything as plain dicts/lists
    (the shape they're already in across the app), adapts them into the
    lightweight objects assemble_proposal_content()/export_proposal_to_docx()
    expect, and returns the finished .docx as bytes — ready for
    st.download_button, same as before.

    deliverables / dates_budget are optional and new: pass
    analysis.get("deliverables", []) / analysis.get("keyDatesBudget", {}) to
    get the populated Key Deliverables table and submission-deadline line
    on the cover page. Omitting them just skips those two pieces — nothing
    else in the export breaks.

    source_label is accepted for backward compatibility with the previous
    version of this function but isn't used in the exported document itself.
    """
    outline_obj = _outline_from_dict(outline)
    responses_obj = _responses_from_list(generated_responses)
    deliverables_obj = _deliverables_from_list(deliverables)
    dates_budget_obj = _dates_budget_from_dict(dates_budget)

    assembled_sections = assemble_proposal_content(
        outline=outline_obj,
        generated_responses=responses_obj,
        rfp_identifier=rfp_identifier,
        company_name=company_name,
    )

    return export_proposal_to_docx(
        rfp_identifier=rfp_identifier,
        assembled_sections=assembled_sections,
        deliverables=deliverables_obj,
        dates_budget=dates_budget_obj,
        company_name=company_name,
    )