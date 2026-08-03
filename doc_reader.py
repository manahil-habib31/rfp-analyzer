"""
doc_reader.py

Multi-format reader for RFP documents. Previously only .pdf was accepted;
this adds .docx, .csv, and .txt, all normalized into the SAME combined-text
format ai_engine.py already expects — one blob with a
"--- Document: filename ---" marker before each file's content. That
convention is what lets ai_engine.py cite which source document a
deliverable/compliance item/question came from (docRef), so every format
here produces text in that same shape rather than each format inventing its
own layout.

PDF extraction itself is NOT reimplemented here — it reuses
pdf_reader.extract_text_from_pdf() exactly as before (including whatever
internal page-marker convention it already produces), so existing PDF
behavior is unchanged. This module only adds the NEW formats and the
combining step across a mixed batch of file types.

DEPENDENCIES: docx (python-docx) is already a dependency of this project
(proposal_builder.py uses it) — no new package needed. CSV/TXT use only the
Python standard library.
"""

import csv
import io
from typing import List, Tuple

from docx import Document as DocxDocument

from pdf_reader import extract_text_from_pdf, PDFExtractionError

SUPPORTED_EXTENSIONS = ["pdf", "docx", "csv", "txt"]


class DocumentExtractionError(Exception):
    """Raised for any non-PDF extraction failure (unreadable .docx, empty
    .csv/.txt, unsupported extension, bad text encoding). Kept as a
    separate class from PDFExtractionError (pdf_reader.py's own exception)
    so callers can catch both without this module needing to know
    PDFExtractionError's internals — but PDF failures are re-raised as this
    same type here too, so app.py only needs to catch ONE exception type at
    the combined-extraction call site."""
    pass


def _decode_bytes(raw: bytes, filename: str) -> str:
    """Text files can arrive as UTF-8 (with or without a BOM) or, less
    commonly, Latin-1 — try UTF-8 first since it's overwhelmingly the norm,
    fall back once, and only then give up with a clear error naming the
    file, rather than crashing with a raw UnicodeDecodeError traceback."""
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentExtractionError(f"Couldn't decode the text encoding of '{filename}'.")


def _extract_docx(file, filename: str) -> str:
    try:
        doc = DocxDocument(file)
    except Exception as e:
        raise DocumentExtractionError(f"Couldn't read '{filename}' as a Word document: {e}") from e

    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Tables often carry real requirements (e.g. a compliance matrix) —
    # skipping them would silently drop content that matters just as much
    # as the paragraph text.
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)

    text = "\n".join(parts)
    if not text.strip():
        raise DocumentExtractionError(f"'{filename}' has no readable text (it may be empty or image-only).")
    return text


def _extract_csv(file, filename: str) -> str:
    raw = file.read()
    text_raw = _decode_bytes(raw, filename) if isinstance(raw, bytes) else raw

    rows = list(csv.reader(io.StringIO(text_raw)))
    if not rows:
        raise DocumentExtractionError(f"'{filename}' appears to be empty.")

    # Render as a simple pipe-delimited table (header + separator + rows)
    # rather than a flat comma blob — keeps column meaning legible to the
    # AI instead of losing which value belongs to which column.
    header, body = rows[0], rows[1:]
    lines = [" | ".join(header), " | ".join(["---"] * len(header))]
    lines.extend(" | ".join(row) for row in body)
    return "\n".join(lines)


def _extract_txt(file, filename: str) -> str:
    raw = file.read()
    text = _decode_bytes(raw, filename) if isinstance(raw, bytes) else raw
    if not text.strip():
        raise DocumentExtractionError(f"'{filename}' appears to be empty.")
    return text


_EXTRACTORS = {
    "docx": _extract_docx,
    "csv": _extract_csv,
    "txt": _extract_txt,
}


def extract_text_from_documents(uploaded_files: List) -> Tuple[str, List[str]]:
    """
    Combines any mix of .pdf, .docx, .csv, and .txt files into ONE text blob
    with a "--- Document: filename ---" marker before each file's content —
    same convention as before, now format-agnostic. Returns
    (combined_text, list_of_filenames_used).

    Raises DocumentExtractionError naming the specific file that failed,
    rather than letting one bad file in a multi-file batch crash the whole
    upload silently.
    """
    combined_parts = []
    doc_names = []

    for f in uploaded_files:
        name = f.name
        ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""

        f.seek(0)
        if ext == "pdf":
            try:
                text = extract_text_from_pdf(f)
            except PDFExtractionError as e:
                raise DocumentExtractionError(f"Couldn't read '{name}': {e}") from e
        elif ext in _EXTRACTORS:
            text = _EXTRACTORS[ext](f, name)
        else:
            raise DocumentExtractionError(
                f"Unsupported file type '.{ext}' for '{name}'. Supported types: "
                + ", ".join(e.upper() for e in SUPPORTED_EXTENSIONS)
            )

        combined_parts.append(f"--- Document: {name} ---\n{text}")
        doc_names.append(name)

    if not combined_parts:
        raise DocumentExtractionError("No documents were provided.")

    return "\n\n".join(combined_parts), doc_names
